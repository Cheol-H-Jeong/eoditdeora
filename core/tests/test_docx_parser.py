from pathlib import Path

from eoditdeora.parsers.docx_parser import DocxParser

from .fixtures import make_docx


def test_docx_roundtrip_paragraphs_and_headings(tmp_path: Path):
    path = make_docx(
        tmp_path / "sample.docx",
        [
            ("Title", "예산 증액 품의서"),
            ("Heading 1", "1. 배경"),
            ("Normal", "전년도 실적을 근거로 예산 증액을 요청합니다."),
            ("Heading 2", "2. 세부 내역"),
            ("Normal", "세부 항목은 첨부 문서를 참조합니다."),
        ],
    )
    res = DocxParser().parse(path, doc_id="sha256:" + "1" * 64)
    assert res.doc.format == "docx"
    assert res.doc.parser == "docx_python_docx"
    heading_levels = [b.level for b in res.doc.blocks if b.type == "heading"]
    # Title treated as level 1, Heading 1 as level 1, Heading 2 as level 2
    assert sorted(heading_levels) == [1, 1, 2]
    paras = [b.text for b in res.doc.blocks if b.type == "paragraph"]
    assert any("전년도" in p for p in paras)


def test_docx_table_extraction(tmp_path: Path):
    import docx

    path = tmp_path / "with_table.docx"
    d = docx.Document()
    d.add_paragraph("문서 본문", style="Normal")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "항목"
    t.cell(0, 1).text = "금액"
    t.cell(1, 0).text = "사무용품"
    t.cell(1, 1).text = "50000"
    d.save(str(path))

    res = DocxParser().parse(path, doc_id="sha256:" + "2" * 64)
    tables = [b for b in res.doc.blocks if b.type == "table"]
    assert len(tables) == 1
    assert tables[0].meta["rows"] == 2
    assert tables[0].meta["cols"] == 2
    assert "사무용품" in tables[0].text
    assert "50000" in tables[0].text


def test_docx_metadata_captured(tmp_path: Path):
    import docx

    path = tmp_path / "meta.docx"
    d = docx.Document()
    d.core_properties.author = "김철수"
    d.core_properties.title = "예산안"
    d.add_paragraph("본문")
    d.save(str(path))

    res = DocxParser().parse(path, doc_id="sha256:" + "3" * 64)
    assert res.doc.metadata.get("author") == "김철수"
    assert res.doc.metadata.get("title") == "예산안"


def test_docx_can_parse_suffix():
    p = DocxParser()
    assert p.can_parse(Path("a.docx"))
    assert p.can_parse(Path("a.DOCX"))
    assert not p.can_parse(Path("a.pdf"))
