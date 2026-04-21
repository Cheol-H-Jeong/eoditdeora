"""DOCX parser via python-docx."""

from __future__ import annotations

import zipfile
from pathlib import Path

from eoditdeora.parsers.base import Block, ParsedDoc, ParseResult
from eoditdeora.parsers.registry import register
from eoditdeora.utils.paths_util import display_path


class DocxParser:
    name = "docx_python_docx"
    supported_extensions = ("docx", "docm")
    fidelity = 5

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower().lstrip(".") in self.supported_extensions

    def parse(self, path: Path, *, doc_id: str) -> ParseResult:
        if not path.exists():
            return ParseResult(
                doc=ParsedDoc(
                    doc_id=doc_id,
                    source_path=str(path),
                    source_path_display=display_path(path),
                    format="docx",
                    parser=self.name,
                    fidelity=self.fidelity,
                    parse_status="file_missing",
                    warnings=["file_missing"],
                )
            )
        if path.stat().st_size == 0:
            return ParseResult(
                doc=ParsedDoc(
                    doc_id=doc_id,
                    source_path=str(path),
                    source_path_display=display_path(path),
                    format="docx",
                    parser=self.name,
                    fidelity=self.fidelity,
                    parse_status="empty",
                    warnings=["empty_file"],
                )
            )

        try:
            import docx  # type: ignore[import-not-found]
        except ImportError as e:
            return ParseResult(
                doc=ParsedDoc(
                    doc_id=doc_id,
                    source_path=str(path),
                    source_path_display=display_path(path),
                    format="docx",
                    parser=self.name,
                    fidelity=self.fidelity,
                    parse_status="parser_error",
                    warnings=[f"python_docx_not_available: {e}"],
                )
            )

        try:
            document = docx.Document(str(path))
        except zipfile.BadZipFile as e:
            return ParseResult(
                doc=ParsedDoc(
                    doc_id=doc_id,
                    source_path=str(path),
                    source_path_display=display_path(path),
                    format="docx",
                    parser=self.name,
                    fidelity=self.fidelity,
                    parse_status="invalid_format",
                    warnings=[f"docx_not_zip: {e}"],
                )
            )
        except Exception as e:  # noqa: BLE001
            return ParseResult(
                doc=ParsedDoc(
                    doc_id=doc_id,
                    source_path=str(path),
                    source_path_display=display_path(path),
                    format="docx",
                    parser=self.name,
                    fidelity=self.fidelity,
                    parse_status="invalid_format",
                    warnings=[f"docx_open_failed: {e}"],
                )
            )

        blocks: list[Block] = []
        for para in document.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading"):
                level = _heading_level(style_name)
                blocks.append(Block(type="heading", text=text, level=level))
            elif style_name in {"title"}:
                blocks.append(Block(type="heading", text=text, level=1))
            elif style_name.startswith("list"):
                blocks.append(Block(type="list_item", text=text))
            elif style_name in {"quote", "intense quote"}:
                blocks.append(Block(type="quote", text=text))
            else:
                blocks.append(Block(type="paragraph", text=text))

        for table in document.tables:
            rows_text: list[list[str]] = []
            for row in table.rows:
                rows_text.append([(c.text or "").strip() for c in row.cells])
            normalized = "\n".join("\t".join(r) for r in rows_text if any(c for c in r))
            if normalized:
                blocks.append(
                    Block(
                        type="table",
                        text=normalized,
                        meta={"rows": len(rows_text), "cols": len(rows_text[0]) if rows_text else 0},
                    )
                )

        core = document.core_properties
        metadata = {
            "author": core.author,
            "created_at": core.created.isoformat() if core.created else None,
            "modified_at": core.modified.isoformat() if core.modified else None,
            "application": "Microsoft Word",
            "title": core.title,
            "subject": core.subject,
        }

        doc = ParsedDoc(
            doc_id=doc_id,
            source_path=str(path),
            source_path_display=display_path(path),
            format="docx",
            parser=self.name,
            fidelity=self.fidelity,
            blocks=blocks,
            metadata={k: v for k, v in metadata.items() if v is not None},
        )
        return ParseResult(doc=doc)


def _heading_level(style_name: str) -> int:
    digits = "".join(c for c in style_name if c.isdigit())
    if not digits:
        return 1
    try:
        return max(1, min(6, int(digits)))
    except ValueError:
        return 1


register(DocxParser())
